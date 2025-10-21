"""
INFORME ECONOMICO DEL EDIFICIO
"""

#!/usr/bin/env python
# coding: utf-8

# In[23]:


import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

def cargar_datos_excel(ruta_archivo, hoja):
    """Función para cargar datos desde un archivo Excel."""
    try:
        with pd.ExcelFile(ruta_archivo) as xls:
            return pd.read_excel(xls, sheet_name=hoja)
    except FileNotFoundError:
        print(f"El archivo no se encontró en la ruta: {ruta_archivo}")
        return None

def filtrar_departamentos_en_mora(df):
    """Función para filtrar departamentos que están en mora."""
    return df[df['MORA'] == 'ESTA EN MORA']

def generar_grafico_barras(df_filtrado):
    """Función para generar un gráfico de barras horizontales con los valores de saldo y mostrar el monto a la izquierda de la barra."""
    plt.figure(figsize=(10, 8))
    barras = plt.barh(df_filtrado['DEPARTAMENTO'], df_filtrado['SALDO'], color='red')
    plt.xlabel('Saldo')
    plt.ylabel('Departamento')
    plt.title('Departamentos en Mora y su saldo deudor')

    # Agregar etiquetas de monto a la izquierda de cada barra
    for barra in barras:
        plt.text(barra.get_width() - 20,  # Colocar el texto a la izquierda de la barra
                 barra.get_y() + barra.get_height() / 2, 
                 f'{barra.get_width():,.2f}', va='center', ha='right')

    plt.tight_layout()

    # Guardar el gráfico como imagen
    ruta_imagen = "grafico_mora.png"
    plt.savefig(ruta_imagen)
    plt.close()
    return ruta_imagen

def crear_informe_word(ruta_imagen, ruta_salida):
    """Función para crear un documento Word e insertar el gráfico."""
    doc = Document()
    doc.add_heading('DEPARTAMENTOS/TIENDAS EN MORA', 0)

    # Insertar el gráfico en el documento Word
    doc.add_picture(ruta_imagen, width=Inches(6))  # Ajustar el tamaño del gráfico

    # Guardar el documento Word
    doc.save(ruta_salida)
    print(f"Informe generado y guardado en: {ruta_salida}")

# Ruta del archivo y hoja específica
file_path = r"C:\Users\HP\Desktop\EDIFICIO JUAN BOSCO\EXPENSAS\EXPENSAS CON CONCILIACION BANCARIA\CONCILIACION BACNCARIA PRUEBAS PYTHON.xlsx"
sheet_name = "CONCILIACION"

# Obtener la carpeta donde está el archivo Excel
carpeta_archivo = os.path.dirname(file_path)
ruta_salida_word = os.path.join(carpeta_archivo, "informe_departamentos_mora.docx")

# Ejecutar el flujo
df = cargar_datos_excel(file_path, sheet_name)
if df is not None:
    df_mora = filtrar_departamentos_en_mora(df)
    if not df_mora.empty:
        ruta_imagen = generar_grafico_barras(df_mora)
        crear_informe_word(ruta_imagen, ruta_salida_word)

        # Eliminar la imagen temporal después de generar el informe
        if os.path.exists(ruta_imagen):
            os.remove(ruta_imagen)
    else:
        print("No se encontraron departamentos en mora.")


# In[ ]:





# In[ ]:






if __name__ == "__main__":
    pass
